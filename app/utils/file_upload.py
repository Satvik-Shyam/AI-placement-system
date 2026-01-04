"""
File Upload Utility - Extract text from resume files.

Supported formats:
- PDF (.pdf) using PyPDF2
- Word (.docx) using python-docx  
- Plain Text (.txt)

Max file size: 5MB
"""

import io
from typing import Tuple
from fastapi import UploadFile, HTTPException

# PDF support
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORTED = True
except ImportError:
    PDF_SUPPORTED = False

# DOCX support
try:
    from docx import Document
    DOCX_SUPPORTED = True
except ImportError:
    DOCX_SUPPORTED = False


MAX_FILE_SIZE_MB = 5
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}


def get_file_extension(filename: str) -> str:
    """Get lowercase file extension."""
    if '.' not in filename:
        return ''
    return '.' + filename.rsplit('.', 1)[1].lower()


async def extract_text_from_file(file: UploadFile) -> Tuple[str, str]:
    """
    Extract text from uploaded file.
    
    Args:
        file: FastAPI UploadFile
        
    Returns:
        Tuple of (extracted_text, filename)
        
    Raises:
        HTTPException on validation/extraction errors
    """
    # Validate filename
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    ext = get_file_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type '{ext}'. Allowed: PDF, DOCX, TXT"
        )
    
    # Read content
    content = await file.read()
    
    # Check size
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size: {MAX_FILE_SIZE_MB}MB"
        )
    
    # Extract based on type
    if ext == '.pdf':
        text = extract_from_pdf(content)
    elif ext == '.docx':
        text = extract_from_docx(content)
    else:  # .txt
        text = extract_from_txt(content)
    
    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from file. File may be empty or corrupted."
        )
    
    return text, file.filename


def extract_from_pdf(content: bytes) -> str:
    """Extract text from PDF bytes."""
    if not PDF_SUPPORTED:
        raise HTTPException(
            status_code=500,
            detail="PDF support not available. Install PyPDF2: pip install PyPDF2"
        )
    
    try:
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return '\n'.join(text_parts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")


def extract_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes."""
    if not DOCX_SUPPORTED:
        raise HTTPException(
            status_code=500,
            detail="DOCX support not available. Install python-docx: pip install python-docx"
        )
    
    try:
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n'.join(text_parts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")


def extract_from_txt(content: bytes) -> str:
    """Extract text from TXT bytes."""
    for encoding in ['utf-8', 'latin-1', 'cp1252']:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="Could not decode text file")


def get_supported_formats() -> dict:
    """Get info about supported file formats."""
    return {
        "supported_formats": [
            {"extension": ".pdf", "available": PDF_SUPPORTED, "name": "PDF"},
            {"extension": ".docx", "available": DOCX_SUPPORTED, "name": "Word Document"},
            {"extension": ".txt", "available": True, "name": "Plain Text"}
        ],
        "max_size_mb": MAX_FILE_SIZE_MB
    }